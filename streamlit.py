import streamlit as st
import whisper
import sounddevice as sd
import queue
import tempfile
import os
import time
import numpy as np
from scipy.io.wavfile import write
from docx import Document
from datetime import datetime

# Load Whisper model once
@st.cache_resource
def load_model():
    return whisper.load_model("base")

model = load_model()

# Audio stream config
samplerate = 16000
channels = 1
duration = 3  # seconds per chunk
q = queue.Queue()

# Streamlit UI
st.set_page_config(page_title="Live Voice Transcriber", layout="centered")
st.title("🎧 Live Voice Transcriber")

# State variables
if "transcript_log" not in st.session_state:
    st.session_state.transcript_log = ""
if "running" not in st.session_state:
    st.session_state.running = False
if "doc" not in st.session_state:
    st.session_state.doc = Document()
    st.session_state.doc.add_heading("Live Voice Transcription", 0)

placeholder = st.empty()
transcript_display = st.empty()

start_btn = st.button("▶ Start Transcribing")
stop_btn = st.button("⏹ Stop")

def callback(indata, frames, time_, status):
    if status:
        st.warning(f"🎤 Mic warning: {status}")
    q.put(indata.copy())

# Start logic
if start_btn:
    st.session_state.running = True
    st.info("🎙️ Listening... Speak now!")
    try:
        with sd.InputStream(samplerate=samplerate, channels=channels, callback=callback):
            while st.session_state.running:
                audio_data = []
                start_time = time.time()
                while time.time() - start_time < duration:
                    audio_data.append(q.get())

                audio_np = np.concatenate(audio_data, axis=0)
                with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as f:
                    write(f.name, samplerate, (audio_np * 32767).astype(np.int16))
                    audio_path = f.name

                result = model.transcribe(audio_path, language="en")  # Always use English letters
                os.remove(audio_path)

                transcript = result['text'].strip()
                if transcript:
                    timestamp = datetime.now().strftime("%H:%M:%S")
                    line = f"[{timestamp}] {transcript}"
                    st.session_state.transcript_log += line + "\n"
                    st.session_state.doc.add_paragraph(line)
                    transcript_display.text_area("Live Transcript", st.session_state.transcript_log, height=400)
    except Exception as e:
        st.error(f"💥 Error: {e}")

# Stop logic
if stop_btn and st.session_state.running:
    st.session_state.running = False
    st.success("🛑 Transcription stopped.")
    # Save transcript
    doc_path = "Live_Transcript.docx"
    st.session_state.doc.save(doc_path)
    with open(doc_path, "rb") as doc_file:
        st.download_button("📄 Download Transcript", doc_file, file_name="Live_Transcript.docx")
